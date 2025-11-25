import os
import logging

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)

BASE_DIR = "clustered_images"
OUTPUT_FILE = "공사사진대지.docx"

# 새 문서 생성
doc = Document()

for cluster in sorted(os.listdir(BASE_DIR)):
    cluster_path = os.path.join(BASE_DIR, cluster)
    images = ["전.jpg", "중.jpg", "후.jpg"]  # 순서대로 존재한다고 가정

    # 표 생성 (4행 2열: 제목 행 1개 + 이미지 3장)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"

    # 제목 셀 병합
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = f"사진대지 - {cluster}"

    # 행 높이 및 텍스트 정렬
    for row in table.rows:
        row.height = Inches(2)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    states = ["전", "중", "후"]
    for i, state in enumerate(states):
        # 상태 텍스트 (전, 중, 후)
        table.cell(i + 1, 0).text = state
        para = table.cell(i + 1, 0).paragraphs[0]
        run = para.runs[0]
        run.font.size = Pt(14)
        para.alignment = 1  # 가운데 정렬

        # 이미지 삽입
        img_path = os.path.join(cluster_path, images[i])
        if os.path.exists(img_path):
            cell = table.cell(i + 1, 1)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(img_path, width=Inches(4.5))

    doc.add_page_break()  # 각 클러스터별로 페이지 구분

# 저장
doc.save(OUTPUT_FILE)
logger.info(f"{OUTPUT_FILE} 생성 완료.")