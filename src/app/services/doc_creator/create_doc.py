import os
import logging

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)


def create_report_docx(base_dir: str, output_file: str):
    """
    Generates a .docx report from clustered image directories.
    """
    if not os.path.exists(base_dir):
        logger.error(f"Base directory for report creation does not exist: {base_dir}")
        raise FileNotFoundError(f"Directory not found: {base_dir}")

    doc = Document()

    cluster_dirs = sorted(
        [d for d in os.listdir(base_dir) if os.path.isdir(os.path.join(base_dir, d))]
    )

    for cluster in cluster_dirs:
        cluster_path = os.path.join(base_dir, cluster)

        images = sorted(
            [
                f
                for f in os.listdir(cluster_path)
                if f.lower().endswith((".png", ".jpg", ".jpeg"))
            ]
        )

        if not images:
            continue

        image_map = {
            "전": images[0] if len(images) > 0 else None,
            "중": images[1] if len(images) > 1 else None,
            "후": images[2] if len(images) > 2 else None,
        }

        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"

        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 0).text = f"사진대지 - {cluster}"

        for row in table.rows:
            row.height = Inches(2)
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        states = ["전", "중", "후"]
        for i, state in enumerate(states):
            table.cell(i + 1, 0).text = state
            para = table.cell(i + 1, 0).paragraphs[0]
            run = para.runs[0]
            run.font.size = Pt(14)
            para.alignment = 1  # 가운데 정렬

            img_name = image_map.get(state)
            if img_name:
                img_path = os.path.join(cluster_path, img_name)
                if os.path.exists(img_path):
                    cell = table.cell(i + 1, 1)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    try:
                        run.add_picture(img_path, width=Inches(4.5))
                    except Exception as e:
                        logger.warning(
                            f"Could not add picture {img_path} to doc: {e}"
                        )

        doc.add_page_break()

    doc.save(output_file)
    logger.info(f"{output_file} 생성 완료.")